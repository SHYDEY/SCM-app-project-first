from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path('SCM_월간_기기·옵션_발주프로세스_구체화_수정본.docx')
NAVY = '1F4D78'
BLUE = '2E74B5'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
GOLD = 'FFF4CC'
RED = 'FCE8E6'
GREEN = 'EAF4EA'
GRAY = '666666'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_font(run, name='Calibri', size=10.5, color='000000', bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def style_paragraph(p, before=0, after=6, line=1.25, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align

def add_para(doc, text='', size=10.5, color='000000', bold=False, italic=False, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, line, align)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    style_paragraph(p, 0, 4, 1.25)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    style_paragraph(p, 0, 4, 1.25)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_font(r, size={1:16,2:13,3:12}[level], color=BLUE if level<3 else NAVY, bold=True)
    return p

def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, 0, 0, 1.08, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            p = cell.paragraphs[0]
            style_paragraph(p, 0, 0, 1.1, WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(val))
            set_font(r, size=font_size)
    return table

def add_callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0,0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, 2, 2, 1.2)
    r = p.add_run(label + '  ')
    set_font(r, size=10.5, color=NAVY, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.5)
    add_para(doc, '', after=2)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Page ')
    set_font(r, size=9, color=GRAY)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.75); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for nm, sz, col, before, after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,NAVY,10,5)]:
    st = styles[nm]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(sz); st.font.color.rgb=RGBColor.from_string(col); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.line_spacing=1.15

header = sec.header.paragraphs[0]
style_paragraph(header, 0, 0, 1.0)
r = header.add_run('한국후지필름BI | SCM 수급 개선 프로젝트')
set_font(r, size=8.5, color=GRAY)
footer = sec.footer.paragraphs[0]
add_page_number(footer)

# Opening block
add_para(doc, 'SCM 수급 개선 프로젝트', size=10, color=BLUE, bold=True, after=2)
p = doc.add_paragraph(); style_paragraph(p, 0, 4, 1.0)
r = p.add_run('월간 기기·옵션 발주 프로세스')
set_font(r, size=24, color=NAVY, bold=True)
add_para(doc, '워크숍 결과 반영 구체화 수정본', size=13, color=GRAY, after=14)
add_table(doc, ['구분','내용'], [
    ['목적','현행 월간 발주 업무를 시스템 요구사항으로 전환할 수 있도록 단계, 판단, 데이터, 예외, PoC 범위를 구체화'],
    ['기준 자료','2차시 프로세스분석 워크숍시트 (2026-08-13) 및 기존 프로세스 초안'],
    ['적용 범위','기기·PRT·옵션 중심의 수요 확정부터 발주 후 입고·PO Match·지불까지'],
    ['핵심 원칙','계산 자동화와 사람의 승인 판단을 분리하고, 모든 수량 조정 사유와 버전을 이력으로 남김'],
], [1450,7910], header_fill=LIGHT_BLUE, font_size=9.6)
add_para(doc, '', after=2)
add_callout(doc, '핵심 정의', '이 시스템은 단순 발주량 계산기가 아니라, OL·SFDC·과거 실적·수급회의의 수요 근거와 재고·Open PO·공급 조건을 연결하여 발주량을 계산하고, 승인·발주·입고·지불까지 추적하는 월간 수급 운영 시스템이다.', fill=GOLD)

heading(doc, '1. 프로세스 범위와 운영 구조', 1)
add_para(doc, '워크숍 시트 기준으로 기존 6단계에 발주 후 관리 단계를 추가하여 총 7단계로 정의한다. 각 단계는 입력물, 담당자, 산출물, 판단 및 시스템 상태를 갖는다.')
add_table(doc, ['단계','업무 목적','주요 산출물','시스템 상태'], [
    ['0','OL 취합 및 발주 사이클 시작','영업부서별 OL 취합본','작성 중'],
    ['①','수요 확인 및 당월 수요 확정','수요 확정본·차이 사유·Bulk-deal 반영 결과','수요 확정'],
    ['②','전월말 재고 및 Open PO 확인','가용재고·미입고 발주 현황','재고 확인'],
    ['③','기기 발주량 산출','기기별 계산 발주량·MOQ/Flex 검증 결과','기기 산출 완료'],
    ['④','옵션 수량 산출','BOM 전개·장착율·Common품 통합 결과','옵션 산출 완료'],
    ['⑤','보고 자료 작성 및 승인','사장 보고 자료·승인/반려 이력','승인 요청/완료'],
    ['⑥','FX-LIVE 등 System 입력','확정 발주 데이터·입력 결과','발주 입력 완료'],
    ['⑦','발주 후 입고·지불 관리','선적/통관/입고/PO Match/지불 상태','입고/지불 관리'],
], [650,2750,3950,2010], font_size=9.1)

heading(doc, '2. 상세 업무 프로세스', 1)
heading(doc, '2.1 0단계: OL 취합 및 사이클 시작', 2)
add_para(doc, '기존 원문에는 명시되지 않았지만 워크숍에서 별도 사이클로 식별된 단계다. 각 영업부서의 OL을 기종·옵션·필요월도 기준으로 취합한 뒤 발주 사이클을 시작한다.')
add_table(doc, ['항목','구체화 내용'], [
    ['시점','월중 발주 마감일을 기준으로 역산. OL 제출 마감일과 지연 시 조치 기한을 별도 설정'],
    ['담당','영업부서: OL 제출 / 수급 담당: 취합·누락 확인 / 물류·구매: 발주 가능성 검토'],
    ['입력','영업부서별 OL, 기종·옵션 코드, 수량, 필요월도, Bulk-deal 표시, 신규·변경·취소 구분'],
    ['산출','통합 OL 버전, 미제출 부서 목록, 코드 오류·중복·수량 이상 목록'],
    ['시스템 요구','업로드 템플릿, 필수 컬럼 검증, 버전 생성, 업로드자·업로드일시 기록'],
], [1500,7860], font_size=9.2)

heading(doc, '2.2 ① 수요 확인 및 당월 수요 확정', 2)
add_para(doc, '수요는 OL을 기본으로 하되 SFDC 중요 파이프라인, OL 외 Bulk-deal, 과거 월도별 실적 Trend를 함께 검토하고 수급회의에서 최종 확정한다.')
add_table(doc, ['세부 업무','입력·확인 내용','산출물·통제'], [
    ['OL 자료 검증','기종·옵션 코드, 필요월도, 수량, 신규/변경/취소, 중복·오기입','오류 목록 및 수정 전후 이력'],
    ['SFDC 중요 파이프라인 확인','기회 단계, 예상 수량, 예상 출고월, 수주 확도, Bulk-deal 여부','반영·참고·제외 구분'],
    ['OL 외 Bulk-deal 확인','대형 프로젝트, 일회성 계약, 사전 재고 확보 필요 여부','별도 수요 라인과 승인 사유'],
    ['과거 Trend 비교','기종별 최근 3·6·12·24개월 실적, 전년 동월, 특수 실적 여부','OL 대비 증감 및 이상치'],
    ['수급회의','수요 출처별 차이, 재고 선확보 필요 여부, 파이프라인 반영 수준','회의 결정사항·참석자·결정일'],
    ['당월 수요 확정','확정 OL + 승인된 조정 - 취소/감소 + 반영 Bulk-deal','수요 버전 및 확정 수량'],
], [1900,4300,3160], font_size=9.0)
add_callout(doc, '확정 원칙', 'SFDC 파이프라인은 확도별 반영률을 정책으로 정하거나 수급회의에서 개별 승인한다. 과거 Trend와 OL이 다를 때는 차이를 숨기지 않고, 최종 선택 근거를 수요 버전에 기록한다.', fill=GOLD)

heading(doc, '2.3 ② 전월말 재고 및 Open PO 확인', 2)
add_para(doc, '전월말 기기 재고를 조회한 뒤 미입고 발주(Open PO) 중 해당 필요월도에 공급 가능한 수량을 추가로 반영하여 가용재고를 계산한다.')
add_table(doc, ['구분','시스템 처리 기준'], [
    ['기말 보유재고','전월 말 마감 기준. 창고·법인·제품 상태별로 조회'],
    ['차감 재고','예약·할당·출고 예정·품질 보류·판매 불가 재고를 정책에 따라 차감'],
    ['Open PO','발주일, 공급사, 품목, 발주 수량, 미입고 수량, 확정 납기, 필요월도별 반영 가능 여부 확인'],
    ['가용재고','보유재고 - 차감재고 + 필요월도 전 입고 가능한 Open PO'],
    ['예외','재고 실사 차이, 입고 지연, 대체 모델, 불용·단종 재고는 별도 예외로 표시'],
], [1900,7460], font_size=9.2)

heading(doc, '2.4 ③ 기기 발주량 산출', 2)
add_para(doc, '기기별 수요·가용재고·Open PO·Lead time·필요월도·Flexibility rule·PRT MOQ를 반영하여 발주량을 계산한다.')
add_table(doc, ['계산 항목','정의·처리 규칙'], [
    ['입고 필요 월도','수요가 실제로 필요한 고객 출고 예측월도. 필요월도와 발주월을 분리 관리'],
    ['Lead time','상해·심천·홍콩·베트남·도쿄·요코하마·나고야·네덜란드 등 Supplier/경로별 기준. 발주~입고 이력으로 실측값 검증'],
    ['기본 필요량','대상 기간 확정 수요 + 목표 안전재고 - 가용재고 - 확정 Open PO'],
    ['Flexibility rule','Supplier 제출 OL 대비 변경 허용 범위 판정. 워크숍 추정값(전월 ±20%, 전전월 ±30%)은 공식 문서 확인 후 확정'],
    ['PRT MOQ','계산량이 MOQ 미달이면 MOQ 단위로 올림하거나 예외 승인. 과잉재고 금액을 보고서에 표시'],
    ['최종 기기 발주량','계산량, MOQ, Flex, 공급 가능량, 담당자 조정량을 반영한 승인 대상 수량'],
], [1900,7460], font_size=9.1)
add_para(doc, '권장 계산식', size=10.5, color=NAVY, bold=True, before=5, after=2)
add_callout(doc, '기본식', '기기 발주 필요량 = 확정 수요 + 목표 안전재고 - 가용 기기 재고 - 필요월도 전 확정 입고량 - 기존 발주 잔량', fill=LIGHT_BLUE)

heading(doc, '2.5 ④ 옵션·PRT 외 필요 수량 산출', 2)
add_para(doc, '옵션은 기기 수요에서 파생되는 종속 수요와 기기와 무관한 A/S·교체 등 별도 수요를 구분한다. BOM과 장착율을 이용해 기종별로 전개한 뒤 Common품을 옵션 코드 기준으로 통합한다.')
add_table(doc, ['세부 업무','구체화 내용','통제 포인트'], [
    ['기기별 필수 옵션 식별','BOM에서 기종별 필수 옵션과 BOM 수량 전개','필수 옵션 누락 시 발주 진행 차단 또는 예외 승인'],
    ['장착율 확인','최근 3·6·12개월 또는 24개월 실적 기준. 신제품은 유사 기종·수급회의 승인값 사용','장착율 기준 기간과 근거 저장'],
    ['평균 사용량 확인','소모품·부품별 3·6·12개월 평균 사용량 및 특수 실적 제외 여부','평균 산식과 제외 기준 저장'],
    ['Common품 합산','복수 기종에서 동일 옵션 코드로 발생한 수요를 합산','옵션 코드 중복 발주 방지'],
    ['별도 옵션 수요','A/S 교체, 서비스, 별도 판매 등 기기 수요와 무관한 수요 추가','수요 유형별 구분 필수'],
    ['재고·LT·Flex·MOQ','옵션 전월말 재고, Open PO, 필요월도, Supplier LT, Flex, MOQ/발주단위 반영','MOQ 과잉 및 납기 미충족 경고'],
], [1900,4500,2960], font_size=8.9)
add_callout(doc, '옵션 계산식', '옵션 필요량 = Σ(기종별 확정 기기 수요 × BOM 수량 × 장착율) + 별도 옵션 수요 + 안전재고 - 옵션 가용재고 - 확정 입고량 - 기존 발주 잔량', fill=LIGHT_BLUE)

heading(doc, '2.6 ⑤ 발주 보고 자료 작성 및 승인', 2)
add_para(doc, '보고서는 총액만 보여주는 자료가 아니라 OL·수요·재고·발주량 차이가 왜 발생했는지를 설명하는 의사결정 자료로 구성한다.')
add_table(doc, ['보고 항목','필수 표시 내용'], [
    ['당월 요약','총 발주금액, 기기 금액, 옵션 금액, 품목 수, 주요 공급사'],
    ['전월·전년 동월 비교','발주금액·수량의 증감, 증감률, 변동 상위 기종·옵션'],
    ['OL 대비 비교','전월 제출 OL, 최종 확정 수요, 실제 발주 필요량, 최종 발주량, 차이 사유'],
    ['재고·공급 리스크','재고 부족, MOQ 과잉, Flex 초과, Lead time상 납기 위험, 단종·대체 모델'],
    ['승인·반려 관리','승인자, 승인일, 조정 수량·금액, 반려 사유, 재계산 시작 단계'],
], [2200,7160], font_size=9.2)
add_para(doc, '사장 보고 반려 시에는 반려 사유에 따라 수요 확정, 재고 확인, 계산, 옵션 산출 중 재시작 지점을 지정한다. 전체를 처음부터 다시 계산하지 않도록 버전과 단계별 확정 상태를 관리한다.')

heading(doc, '2.7 ⑥ 발주량 System 입력 및 검증', 2)
add_table(doc, ['입력 단계','필수 처리'], [
    ['발주 데이터 확정','승인 완료 버전만 입력 대상으로 잠금. 담당자 조정값과 승인 이력 포함'],
    ['FX-LIVE 등 System 입력','공급사, 품목 코드, 필요월도, 수량, 단가, 통화, 납기 요청일, 승인번호 입력'],
    ['입력 결과 검증','입력 수량·금액·품목·납기일을 산출 결과와 대조. 오류 메시지와 재처리 이력 저장'],
    ['발주번호 연결','System 발주번호를 산출 버전에 연결하여 이후 선적·입고·PO Match 추적'],
], [2200,7160], font_size=9.2)

heading(doc, '2.8 ⑦ 발주 후 선적·통관·입고·지불 관리', 2)
add_para(doc, '워크숍에서 식별된 발주 후 업무를 시스템 범위에 포함한다. 발주가 끝나는 시점이 아니라 Supplier 접수부터 지불까지가 하나의 발주 건으로 추적되어야 한다.')
add_table(doc, ['순서','업무','관리 데이터'], [
    ['1','Supplier 발주 접수 및 수량 확인','접수일, Supplier 확인 수량, 납기, 변경·취소 여부'],
    ['2','현지 선적 및 일정 관리','선적일, 선적 여부, ETA, 운송장, 지연 사유, 기상·재난·전쟁 리스크'],
    ['3','한국 도착 및 통관 준비','부산항·인천항·인천공항 도착 예정, Invoice·Packing List·B/L 등 통관 서류'],
    ['4','통관 및 창고 입고','통관 완료일, 입고 예정일, 실제 입고일, 창고, 입고 수량'],
    ['5','입고 수량·상태 검수','수량 차이, 파손·불량, 미입고·과입고, 검수 결과'],
    ['6','PO Match 및 지불 처리','PO·입고·Invoice 매칭, 차이 금액, 지급 승인일, Supplier별 지급 상태'],
], [700,3000,5660], font_size=9.1)

heading(doc, '3. 사람이 판단해야 하는 항목과 자동화 경계', 1)
add_para(doc, '자동화의 목표는 판단을 없애는 것이 아니라, 반복 조회·계산·비교는 자동화하고 사업적 판단과 예외 승인은 사람이 수행하도록 경계를 명확히 하는 것이다.')
add_table(doc, ['판단 항목','권장 처리','사람의 역할'], [
    ['당월 수요 확정','수요 출처 통합·차이 자동 제시','수급회의에서 반영·제외·조정 결정'],
    ['파이프라인 반영 수준','확도·정책 기반 후보 수량 계산','Bulk-deal 및 전략 프로젝트 최종 승인'],
    ['Trend 이탈','전년·최근 평균 대비 이상치 경고','특수 실적·신제품·시장 이슈 해석'],
    ['가용재고 기준','예약·할당·보류 재고 자동 분류','정책 예외와 재고 실사 차이 승인'],
    ['Lead time·필요월도','기준정보 및 실측 이력 기반 계산','긴급·변경·공급사 협의 판단'],
    ['Flex·MOQ 위반','자동 판정 및 경고','초과 발주·과잉 재고 감수 여부 승인'],
    ['옵션 장착율·평균 사용량','실적 기반 자동 산출','신제품·특수 프로젝트 적용값 조정'],
    ['보고서 차이 원인','전월·전년·OL 대비 자동 분해','경영진 설명용 핵심 사유 확정'],
], [1900,3650,3810], font_size=8.9)

heading(doc, '4. 예외·누락 구간 및 처리 방향', 1)
add_table(doc, ['예외 상황','시스템 처리 방향','확인 필요 사항'], [
    ['OL 지연·누락','미제출 부서와 미확정 수요를 경고하고 대체 기준 사용 여부 승인','제출 마감일·지연 승인자'],
    ['Flex rule 초과','초과량·초과 사유·Supplier 협의 상태를 표시하고 승인 전송','본사 협의 절차'],
    ['월중 긴급 발주','정규 발주와 별도 유형으로 생성하고 긴급 사유·추가 비용 기록','발생 빈도·PoC 포함 여부'],
    ['확정 발주 변경·취소','기존 발주와 변경 발주를 연결하고 변경 가능 단계 및 비용 기록','Supplier별 변경 가능 시점'],
    ['신제품·단종','유사 기종 기반 장착율 또는 수요를 임시 적용하고 유효기간 관리','신제품 초기 기준·단종 처리'],
    ['재고 실사 차이','실사 조정 전후 재고와 승인자를 기록하고 발주 계산 기준일을 고정','창고 범위·실사 주기'],
    ['A/S 등 별도 옵션 수요','기기 종속 수요와 별도 수요 유형을 분리','옵션 수요 원천·담당자'],
    ['공급·물류 장애','선적·도착·통관 지연을 리스크 상태로 관리하고 수요·대체 모델 영향 표시','모니터링 담당·주기'],
], [2100,4300,2960], font_size=8.8)

heading(doc, '5. 시스템 데이터 요구사항', 1)
add_table(doc, ['데이터','용도','최소 필드','우선순위'], [
    ['기종 마스터','기기 발주 계산','기종코드, 제품군, Supplier, 단가, MOQ, 발주단위, LT, Flex, 단종·대체','높음'],
    ['옵션 마스터','옵션 발주 계산','옵션코드, 단가, MOQ, 발주단위, Common 여부, Supplier, LT','높음'],
    ['BOM','기종-옵션 수요 전개','기종코드, 옵션코드, BOM 수량, 필수 여부, 유효기간','높음'],
    ['기기·옵션 재고','가용재고 계산','기말일, 품목, 창고, 보유·예약·보류·가용 수량','높음'],
    ['Open PO','미입고 수량 반영','PO번호, 발주일, 품목, 발주·미입고 수량, 확정 납기, Supplier','높음'],
    ['기기·옵션 실적 24개월','Trend·장착율·평균 사용량','월도, 품목, 수량, 수요 유형, 고객 식별자 마스킹','중간'],
    ['발주~입고 이력','LT 실측','PO번호, 발주일, 선적일, 도착일, 입고일, 수량','최우선'],
    ['OL 12개월 이력','Flex·OL 대비 분석','월도, 품목, 제출 수량, 수정 이력, 확정 여부','높음'],
    ['SFDC 파이프라인','추가 수요 반영','기회, 단계, 확도, 수량, 필요월도, Bulk-deal 여부','중간'],
    ['보고 샘플·기존 산출 엑셀','현행 로직·출력 형식 확인','파일 버전, 산출식, 보고 항목, 승인 흔적','높음'],
], [1700,2300,4350,1010], font_size=8.4)
add_callout(doc, '최우선 데이터', '발주~입고 이력에는 발주일과 입고일이 같은 발주 건으로 연결되어야 한다. 이 연결이 없으면 Supplier별 Lead time 실측과 납기 예측을 검증할 수 없다.', fill=RED)

heading(doc, '6. PoC 권장 범위', 1)
add_para(doc, '워크숍 시트의 PoC 후보를 실제 구축 순서로 재정렬하면 다음 범위가 1차 검증에 적합하다.')
add_table(doc, ['구분','1차 PoC 포함','후속 또는 제외 후보'], [
    ['프로세스','① 수요 통합·차이 제시, ② 재고·Open PO 집계, ③ 기기 발주량, ④ 옵션 수량, ⑤ 보고 초안, ⑥ 발주 데이터 파일 출력','⑦ 선적·통관·입고·지불 전체 운영 자동화'],
    ['기능','MOQ, Flex 자동 검증, 장착율 계산, Common품 합산, 필수 옵션 누락 경고, 담당자 조정·사유 기록','자연어 AI Assistant는 데이터·규칙 확정 후 별도 검토'],
    ['데이터 연동','초기에는 템플릿 업로드와 파일 출력으로 검증','ERP·SFDC 실시간 연동'],
    ['운영 범위','기기·PRT·옵션 중 대표 기종·옵션으로 산식 검증','월중 긴급 발주, 다월도 롤링, 유지보수 부품'],
], [1700,4400,3260], font_size=8.8)

heading(doc, '7. 시스템 기능 요구사항 요약', 1)
add_table(doc, ['기능 영역','필수 기능'], [
    ['수요 관리','OL 업로드·검증, SFDC·Trend 비교, Bulk-deal 관리, 수급회의 결정사항, 수요 버전·승인'],
    ['재고 관리','전월말 재고, 예약·보류·가용재고, Open PO, 입고 예정, 재고 기준일 고정'],
    ['발주 계산','기기·옵션 산식, LT·필요월도·안전재고, MOQ·발주단위·Flex, 공통품 통합'],
    ['판단·예외','초과·부족·납기·MOQ·필수 옵션·데이터 누락 경고, 조정 사유·승인'],
    ['보고','당월 총액, 전월·전년 동월 비교, OL 대비 차이, 주요 원인·리스크, Excel/PDF 출력'],
    ['승인·입력','금액·예외별 승인선, 반려 재시작 지점, FX-LIVE 입력 파일 또는 연동, 입력 결과 검증'],
    ['사후관리','Supplier 접수, 선적, 도착, 통관, 입고, 검수, PO Match, 지불 상태 추적'],
    ['감사·권한','버전, 변경 이력, 사용자·일시, 기준정보 유효기간, 역할별 권한'],
], [1900,7460], font_size=9.1)

heading(doc, '8. 구축 전 확정해야 할 질문', 1)
add_number(doc, '월간 발주 사이클의 시작일·OL 제출 마감일·수급회의일·사장 보고일·System 입력 마감일은 언제인가?')
add_number(doc, 'OL은 누가 어떤 근거로 제출하며, 지연·누락 시 어떤 대체 기준을 적용하는가?')
add_number(doc, 'Flexibility rule의 공식 문서와 Supplier별·제품군별 차이는 무엇인가? 워크숍의 전월 ±20%, 전전월 ±30% 추정값을 확정할 수 있는가?')
add_number(doc, '안전재고·예약재고·할당재고·품질보류재고를 가용재고 계산에서 어떻게 처리하는가?')
add_number(doc, 'PRT와 그 외 제품군의 MOQ·발주단위·예외 승인 기준이 다른가?')
add_number(doc, '장착율·평균 사용량 산출 기간과 특수·프로젝트 실적 제외 기준은 무엇인가?')
add_number(doc, '기기와 무관한 옵션 수요(A/S·교체 등)의 데이터 원천과 담당자는 누구인가?')
add_number(doc, '발주 후 선적·통관·입고·PO Match·지불을 PoC에서 어느 수준까지 포함할 것인가?')
add_number(doc, '⑥ System은 FX-LIVE가 맞는지, 입력 파일 업로드인지 API 연동인지 확정되었는가?')

heading(doc, '9. 최종 프로세스 요약', 1)
add_callout(doc, '운영 흐름', 'OL 취합 → OL 검증·SFDC·Bulk-deal·Trend 확인 → 수급회의 → 당월 수요 확정 → 전월말 재고·Open PO 확인 → 기기 발주량 계산 → 옵션 BOM 전개·장착율·평균 사용량·Common품·필수품 반영 → LT·필요월도·Flex·MOQ 검증 → 보고·승인 → FX-LIVE 입력 → Supplier 접수·선적·통관·입고·검수·PO Match·지불 → 실적 및 차이 분석', fill=GREEN)
add_para(doc, '이 구조를 기반으로 다음 산출물은 ① 월간 발주 프로세스 정의서, ② 데이터 정의서, ③ 발주량 산식·정책서, ④ 화면·승인·예외 요구사항, ⑤ PoC 테스트 시나리오로 분리하면 된다.', after=0)

# Core properties / metadata
doc.core_properties.title = '월간 기기·옵션 발주 프로세스 구체화 수정본'
doc.core_properties.subject = 'SCM 수급 개선 프로젝트'
doc.core_properties.author = 'SCM 수급 개선 프로젝트'
doc.save(OUT)
print(OUT.resolve())
